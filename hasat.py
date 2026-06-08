#!/usr/bin/env python3
"""
Hasat - Web Wordlist Generator
GitHub: github.com/bolatbatuhan/hasat

Recursively crawls a target website to generate a wordlist from HTML, 
JS-rendered content, meta/alt tags, and file attachments (PDF, DOCX, XLSX). 
Features built-in support for email and username extraction.

# Usage:
python hasat.py -u https://target.local -d 2 -o wordlist.txt
python hasat.py -u https://target.local -d 3 --js --files --email -v
python hasat.py -u https://target.local --proxy http://127.0.0.1:8080 --no-verify-ssl
"""

import re
import sys
import time
import argparse
import tempfile
import os
import requests

from collections import Counter
from urllib.parse import urljoin, urlparse
from requests.auth import HTTPBasicAuth, HTTPDigestAuth

try:
    from playwright.sync_api import sync_playwright
    PLAYWRIGHT_OK = True
except ImportError:
    PLAYWRIGHT_OK = False

try:
    import pdfplumber
    PDFPLUMBER_OK = True
except ImportError:
    PDFPLUMBER_OK = False

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import openpyxl
    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False


# ─────────────────────────────────────────────
#  Terminal Colors
# ─────────────────────────────────────────────
class C:
    RED     = "\033[91m"
    GREEN   = "\033[92m"
    YELLOW  = "\033[93m"
    CYAN    = "\033[96m"
    MAGENTA = "\033[95m"
    RESET   = "\033[0m"
    BOLD    = "\033[1m"
    DIM     = "\033[2m"


def banner():
    print(f"""{C.GREEN}{C.BOLD}
  ██╗  ██╗ █████╗ ███████╗ █████╗ ████████╗
  ██║  ██║██╔══██╗██╔════╝██╔══██╗╚══██╔══╝
  ███████║███████║███████╗███████║   ██║   
  ██╔══██║██╔══██║╚════██║██╔══██║   ██║   
  ██║  ██║██║  ██║███████║██║  ██║   ██║   
  ╚═╝  ╚═╝╚═╝  ╚═╝╚══════╝╚═╝  ╚═╝   ╚═╝   
{C.RESET}{C.YELLOW}  Web Wordlist & OSINT Harvester{C.RESET}
{C.DIM}  ─────────────────────────────────────────{C.RESET}
""")


# ─────────────────────────────────────────────
#  Regex & Text Processing
# ─────────────────────────────────────────────

EMAIL_RE    = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
USERNAME_RE = re.compile(r"[a-zA-Z][a-zA-Z0-9._\-]{2,32}")


def normalize_word(w: str) -> str:
    return re.sub(r"^[^a-zA-Z0-9]+|[^a-zA-Z0-9]+$", "", w).lower()


def extract_words_from_text(text: str, min_len: int, max_len: int) -> list[str]:
    raw = re.findall(r"[a-zA-Z0-9][\w\-]*[a-zA-Z0-9]|[a-zA-Z0-9]+", text)
    out = []
    for w in raw:
        c = normalize_word(w)
        if c and min_len <= len(c) <= max_len:
            out.append(c)
    return out


def extract_emails(text: str) -> set[str]:
    return set(m.lower() for m in EMAIL_RE.findall(text))


def derive_usernames(emails: set[str]) -> set[str]:
    """
    Parse standard corporate email naming conventions to extract usernames
    Target schema: firstname.lastname@target.local -> [firstname.lastname, f.lastname, lastname]
    """
    usernames = set()
    for email in emails:
        local = email.split("@")[0]         
        usernames.add(local)
        parts = re.split(r"[._\-]", local)  
        if len(parts) >= 2:
            usernames.add(parts[0])                                   
            usernames.add(parts[-1])                                  
            usernames.add(parts[0][0] + parts[-1])          
            usernames.add(parts[0] + parts[-1])              
            usernames.add(parts[0] + "." + parts[-1])        
    return usernames


# ─────────────────────────────────────────────
#  HTML Parse: meta, alt, title, body
# ─────────────────────────────────────────────

def parse_html(html: str, min_len: int, max_len: int) -> dict:
    result = {
        "words":  [],
        "emails": set(),
        "meta":   [],
        "alt":    [],
    }

    body_text = re.sub(r"<[^>]+>", " ", html)
    body_text = re.sub(r"&[a-zA-Z]+;", " ", body_text)
    result["words"] += extract_words_from_text(body_text, min_len, max_len)
    result["emails"] |= extract_emails(body_text)

    title_match = re.search(r"<title[^>]*>([^<]+)</title>", html, re.IGNORECASE)
    if title_match:
        result["words"] += extract_words_from_text(title_match.group(1), min_len, max_len)

    meta_contents = re.findall(
        r'<meta[^>]+content=["\']([^"\']+)["\']',
        html, re.IGNORECASE
    )
    for content in meta_contents:
        words = extract_words_from_text(content, min_len, max_len)
        result["words"] += words
        result["meta"]  += words

    alt_contents = re.findall(r'\balt=["\']([^"\']+)["\']', html, re.IGNORECASE)
    for alt in alt_contents:
        words = extract_words_from_text(alt, min_len, max_len)
        result["words"] += words
        result["alt"]   += words

    return result


def extract_links(html: str, base_url: str) -> set[str]:
    base_parsed = urlparse(base_url)
    hrefs = re.findall(r'href=["\']([^"\'#?]+)["\']', html, re.IGNORECASE)
    links = set()
    for href in hrefs:
        full   = urljoin(base_url, href)
        parsed = urlparse(full)
        if parsed.scheme in ("http", "https") and parsed.netloc == base_parsed.netloc:
            clean = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
            if clean.rstrip("/") != base_url.rstrip("/"):
                links.add(clean)
    return links


def extract_file_links(html: str, base_url: str) -> dict[str, list[str]]:
    extensions = {
        "pdf":  [],
        "docx": [],
        "xlsx": [],
    }
    all_hrefs = re.findall(r'href=["\']([^"\']+)["\']', html, re.IGNORECASE)
    for href in all_hrefs:
        full = urljoin(base_url, href)
        lower = full.lower().split("?")[0]
        for ext in extensions:
            if lower.endswith(f".{ext}"):
                extensions[ext].append(full)
    return extensions


# ─────────────────────────────────────────────
#  File Content Extraction
# ─────────────────────────────────────────────

def read_pdf(data: bytes, min_len: int, max_len: int) -> list[str]:
    if not PDFPLUMBER_OK:
        return []
    words = []
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        with pdfplumber.open(tmp_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                words += extract_words_from_text(text, min_len, max_len)
    except Exception:
        pass
    finally:
        os.unlink(tmp_path)
    return words


def read_docx(data: bytes, min_len: int, max_len: int) -> list[str]:
    if not DOCX_OK:
        return []
    words = []
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        doc = DocxDocument(tmp_path)
        for para in doc.paragraphs:
            words += extract_words_from_text(para.text, min_len, max_len)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    words += extract_words_from_text(cell.text, min_len, max_len)
    except Exception:
        pass
    finally:
        os.unlink(tmp_path)
    return words


def read_xlsx(data: bytes, min_len: int, max_len: int) -> list[str]:
    if not OPENPYXL_OK:
        return []
    words = []
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        wb = openpyxl.load_workbook(tmp_path, read_only=True, data_only=True)
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell and isinstance(cell, str):
                        words += extract_words_from_text(cell, min_len, max_len)
    except Exception:
        pass
    finally:
        os.unlink(tmp_path)
    return words


# ─────────────────────────────────────────────
#  Session Builder
# ─────────────────────────────────────────────

def build_session(args) -> requests.Session:
    session = requests.Session()
    ua = args.user_agent or (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )
    headers = {"User-Agent": ua}

    if args.header:
        for h in args.header:
            if ":" in h:
                k, v = h.split(":", 1)
                headers[k.strip()] = v.strip()
            else:
                print(f"{C.YELLOW}[!] Invalid header format (skipping): {h}{C.RESET}")

    session.headers.update(headers)

    if args.auth_user and args.auth_pass:
        if args.auth_type == "digest":
            session.auth = HTTPDigestAuth(args.auth_user, args.auth_pass)
            print(f"{C.CYAN}[*] Digest Auth: {args.auth_user}{C.RESET}")
        else:
            session.auth = HTTPBasicAuth(args.auth_user, args.auth_pass)
            print(f"{C.CYAN}[*] Basic Auth: {args.auth_user}{C.RESET}")

    if args.proxy:
        session.proxies = {"http": args.proxy, "https": args.proxy}
        print(f"{C.CYAN}[*] Proxy: {args.proxy}{C.RESET}")

    session.verify = not args.no_verify_ssl
    return session


# ─────────────────────────────────────────────
#  Main Crawler: Hasat
# ─────────────────────────────────────────────

class Hasat:
    def __init__(self, args):
        self.start_url    = args.url.rstrip("/")
        self.max_depth    = args.depth
        self.min_len      = args.min_word
        self.max_len      = args.max_word
        self.output_file  = args.output
        self.delay        = args.delay
        self.verbose      = args.verbose
        self.use_js       = args.js
        self.scan_files   = args.files
        self.extract_email= args.email
        self.proxy        = args.proxy
        self.session      = build_session(args)

        self.visited      = set()
        self.fetched_files= set()
        self.word_counter = Counter()
        self.emails       = set()
        self.usernames    = set()

        # Stats
        self.stats = {
            "pages":    0,
            "files":    0,
            "meta_words": 0,
            "alt_words":  0,
            "email_count":0,
        }

    # ── HTTP fetch ──────────────────────────
    def fetch_html(self, url: str) -> str | None:
        try:
            r = self.session.get(url, timeout=12)
            r.raise_for_status()
            return r.text
        except requests.exceptions.SSLError:
            print(f"{C.RED}[!] SSL error: {url}{C.RESET}")
        except requests.exceptions.ProxyError:
            print(f"{C.RED}[!] Proxy error: {url}{C.RESET}")
        except requests.exceptions.ConnectionError:
            print(f"{C.RED}[!] Connection error: {url}{C.RESET}")
        except requests.exceptions.HTTPError as e:
            print(f"{C.RED}[!] HTTP {e.response.status_code}: {url}{C.RESET}")
        except requests.exceptions.Timeout:
            print(f"{C.YELLOW}[!] Timeout: {url}{C.RESET}")
        return None

    def fetch_bytes(self, url: str) -> bytes | None:
        try:
            r = self.session.get(url, timeout=20)
            r.raise_for_status()
            return r.content
        except Exception as e:
            print(f"{C.RED}[!] Failed to download file ({url}): {e}{C.RESET}")
        return None

    # ── JS Render (Playwright) ──────────────
    def fetch_js_rendered(self, url: str) -> str | None:
        if not PLAYWRIGHT_OK:
            print(f"{C.YELLOW}[!] Playwright is not installed. Run: 'pip install playwright && playwright install chromium'{C.RESET}")
            return None
        try:
            with sync_playwright() as p:
                launch_args = {"headless": True}
                if self.proxy:
                    launch_args["proxy"] = {"server": self.proxy}

                browser = p.chromium.launch(**launch_args)
                page    = browser.new_page()
                page.goto(url, timeout=15000, wait_until="networkidle")
                html = page.content()
                browser.close()
                return html
        except Exception as e:
            print(f"{C.RED}[!] JS render error ({url}): {e}{C.RESET}")
            return None

    # ── File Content Extraction ────────────────────────
    def process_files(self, file_links: dict, base_url: str):
        handlers = {
            "pdf":  (read_pdf,  PDFPLUMBER_OK, "pdfplumber"),
            "docx": (read_docx, DOCX_OK,       "python-docx"),
            "xlsx": (read_xlsx, OPENPYXL_OK,   "openpyxl"),
        }
        for ext, urls in file_links.items():
            fn, available, pkg = handlers[ext]
            if not available:
                if urls:
                    print(f"{C.YELLOW}[!] Install '{pkg}' for {ext.upper()} processing support.{C.RESET}")
                continue
            for url in urls:
                if url in self.fetched_files:
                    continue
                self.fetched_files.add(url)
                print(f"  {C.MAGENTA}[file]{C.RESET} {ext.upper()}: {url}")
                data = self.fetch_bytes(url)
                if data:
                    words = fn(data, self.min_len, self.max_len)
                    self.word_counter.update(words)
                    self.stats["files"] += 1
                    if self.verbose:
                        print(f"    → {len(words)} words harvested ({ext})")

    # ── Main crawl ───────────────────────────
    def crawl(self, url: str, depth: int):
        if depth > self.max_depth or url in self.visited:
            return
        self.visited.add(url)
        self.stats["pages"] += 1

        prefix = f"[{depth}/{self.max_depth}]"
        print(f"  {C.GREEN}{prefix}{C.RESET} {url}")

        if self.use_js:
            html = self.fetch_js_rendered(url) or self.fetch_html(url)
        else:
            html = self.fetch_html(url)

        if not html:
            return

        parsed = parse_html(html, self.min_len, self.max_len)
        self.word_counter.update(parsed["words"])
        self.stats["meta_words"] += len(parsed["meta"])
        self.stats["alt_words"]  += len(parsed["alt"])

        if self.extract_email:
            new_emails = parsed["emails"]
            if new_emails:
                self.emails |= new_emails
                usernames    = derive_usernames(new_emails)
                self.usernames |= usernames
                if self.verbose:
                    for em in new_emails:
                        print(f"    {C.CYAN}[email]{C.RESET} {em}")
                    for un in usernames:
                        print(f"    {C.CYAN}[user] {C.RESET} {un}")

        if self.verbose:
            print(f"    → {len(parsed['words'])} words | "
                  f"meta:{len(parsed['meta'])} alt:{len(parsed['alt'])}")

        if self.scan_files:
            file_links = extract_file_links(html, url)
            self.process_files(file_links, url)

        if depth < self.max_depth:
            links = extract_links(html, url)
            for link in sorted(links):
                if link not in self.visited:
                    time.sleep(self.delay)
                    self.crawl(link, depth + 1)

    # ── Save Outputs ──────────────────────────────
    def save_wordlist(self):
        words = sorted(self.word_counter.keys())
        with open(self.output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(words) + "\n")
        print(f"\n  {C.GREEN}✓{C.RESET} Wordlist  → {self.output_file} ({len(words)} words)")

    def save_emails(self):
        if not self.emails:
            return
        email_file = self.output_file.replace(".txt", "") + "_emails.txt"
        with open(email_file, "w", encoding="utf-8") as f:
            for em in sorted(self.emails):
                f.write(em + "\n")
        print(f"  {C.GREEN}✓{C.RESET} Emails    → {email_file} ({len(self.emails)} addresses)")

    def save_usernames(self):
        if not self.usernames:
            return
        user_file = self.output_file.replace(".txt", "") + "_usernames.txt"
        with open(user_file, "w", encoding="utf-8") as f:
            for un in sorted(self.usernames):
                f.write(un + "\n")
        print(f"  {C.GREEN}✓{C.RESET} Usernames → {user_file} ({len(self.usernames)} items)")

    # ── Run ─────────────────────────────────
    def run(self):
        print(f"{C.BOLD}Target   : {self.start_url}{C.RESET}")
        print(f"{C.BOLD}Depth    : {self.max_depth}  |  "
              f"Length: {self.min_len}-{self.max_len} chars  |  "
              f"JS: {'✓' if self.use_js else '✗'}  |  "
              f"Files: {'✓' if self.scan_files else '✗'}  |  "
              f"Email: {'✓' if self.extract_email else '✗'}{C.RESET}\n")

        self.crawl(self.start_url, depth=0)

        print(f"\n{C.BOLD}{'─'*52}{C.RESET}")
        print(f"{C.GREEN}{C.BOLD}  Harvesting completed!{C.RESET}")
        print(f"  Pages crawled   : {self.stats['pages']}")
        print(f"  Files processed : {self.stats['files']}")
        print(f"  Meta words      : {self.stats['meta_words']}")
        print(f"  Alt words       : {self.stats['alt_words']}")
        print(f"  Emails found    : {len(self.emails)}")
        print(f"  Usernames derived: {len(self.usernames)}")
        print(f"  Total unique words: {len(self.word_counter)}")

        if self.verbose and self.word_counter:
            print(f"\n{C.CYAN}  Top 10 most common:{C.RESET}")
            for w, n in self.word_counter.most_common(10):
                print(f"    {w:<28} ({n}x)")

        print()
        self.save_wordlist()
        if self.extract_email:
            self.save_emails()
            self.save_usernames()


# ─────────────────────────────────────────────
#  CLI Argument Parser
# ─────────────────────────────────────────────

def parse_args():
    p = argparse.ArgumentParser(
        prog="hasat",
        description="Hasat — Web Wordlist & OSINT Harvester",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python hasat.py -u https://target.local -d 2 -o wordlist.txt
  python hasat.py -u https://target.local -d 3 --js --files --email -v
  python hasat.py -u https://target.local --auth-user pen-tester --auth-pass 'SecurePass123!' --auth-type digest
  python hasat.py -u https://target.local --proxy http://127.0.0.1:8080 --no-verify-ssl
  python hasat.py -u https://target.local --header "Cookie: session=abc" --header "X-Forwarded-For: 1.2.3.4"
        """
    )

    p.add_argument("-u", "--url",       required=True, help="Target URL")
    p.add_argument("-d", "--depth",      type=int, default=2, help="Crawl depth (default: 2)")
    p.add_argument("-o", "--output",     default="wordlist.txt", help="Output wordlist file (default: wordlist.txt)")
    p.add_argument("-v", "--verbose",    action="store_true", help="Enable verbose output")
    p.add_argument("--delay",            type=float, default=0.5, help="Delay between page requests in seconds (default: 0.5)")

    p.add_argument("--min-word", type=int, default=4,  dest="min_word", help="Minimum word length (default: 4)")
    p.add_argument("--max-word", type=int, default=25, dest="max_word", help="Maximum word length (default: 25)")

    p.add_argument("--js",    action="store_true", help="Enable JavaScript rendering via Playwright")
    p.add_argument("--files", action="store_true", help="Download and extract content from PDF, DOCX, and XLSX files")
    p.add_argument("--email", action="store_true", help="Extract email addresses and derive custom usernames")

    p.add_argument("--auth-user", dest="auth_user", help="HTTP Authentication username")
    p.add_argument("--auth-pass", dest="auth_pass", help="HTTP Authentication password")
    p.add_argument("--auth-type", dest="auth_type", choices=["basic", "digest"], default="basic",
                   help="HTTP Authentication type (default: basic)")

    p.add_argument("--user-agent", dest="user_agent", help="Custom User-Agent header")
    p.add_argument("--header",     action="append",   help="Custom HTTP header: 'Key: Value' (can be used multiple times)")

    p.add_argument("--proxy",          help="Proxy URL (e.g., http://127.0.0.1:8080 or socks5://127.0.0.1:9050)")
    p.add_argument("--no-verify-ssl", action="store_true", dest="no_verify_ssl",
                   help="Disable SSL/TLS certificate verification")

    return p.parse_args()


if __name__ == "__main__":
    banner()
    args = parse_args()

    if not args.url.startswith(("http://", "https://")):
        print(f"{C.RED}[!] URL must start with 'http://' or 'https://'.{C.RESET}")
        sys.exit(1)

    if args.min_word > args.max_word:
        print(f"{C.RED}[!] --min-word cannot be greater than --max-word.{C.RESET}")
        sys.exit(1)

    if args.js and not PLAYWRIGHT_OK:
        print(f"{C.YELLOW}[!] To use --js option, run: pip install playwright && playwright install chromium{C.RESET}")

    try:
        h = Hasat(args)
        h.run()
    except KeyboardInterrupt:
        print(f"\n{C.YELLOW}[!] Execution interrupted by user. Saving gathered data...{C.RESET}")
        try:
            h.save_wordlist()
            if args.email:
                h.save_emails()
                h.save_usernames()
        except Exception:
            pass
        sys.exit(0)
        